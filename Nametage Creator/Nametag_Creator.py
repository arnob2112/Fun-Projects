from docx import Document
from docx2pdf import convert

replacements = {
    "OC": "OC",
    "এহসান": "আহাবাব",
    "EHSHAN": "AHABAB",
    "M-4(A)": "M-4(A)",
    "7.": "4.",
    "14299": "14281"
}

doc = Document("Nametag - Printable - Previous.docx")


def replace_in_paragraphs(paragraphs):
    for paragraph in paragraphs:
        # normal run-based replacement (your original logic)
        for run in paragraph.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)

        # fallback for split-run cases (serial number fix)
        full_text = "".join(run.text for run in paragraph.runs)
        for old, new in replacements.items():
            if old in full_text:
                full_text = full_text.replace(old, new)

        # write back ONLY if needed
        if paragraph.runs and full_text != "".join(run.text for run in paragraph.runs):
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""


# Normal paragraphs
replace_in_paragraphs(doc.paragraphs)

# Table (single table, single row, no nested tables)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            replace_in_paragraphs(cell.paragraphs)


doc.save("Nametag - Printable.docx")

convert("Nametag - Printable.docx", "New - Nametag - Printable.pdf")

print("Your new nametag file is ready.")
